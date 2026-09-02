from pathlib import Path

import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "latitude_zone_analysis"
OUT_DIR = ROOT / "reports" / "article_gost_materials"
OUT_DIR.mkdir(parents=True, exist_ok=True)

ZONE_ORDER = ["Low", "S_mid", "N_mid", "N_high"]
ZONE_LABELS = {
    "Low": "Низкие широты",
    "S_mid": "Южные средние",
    "N_mid": "Северные средние",
    "N_high": "Северные высокие",
}
MODEL_ORDER = ["LinearRegression", "ElasticNet", "RandomForest", "XGBoost", "CatBoost"]
COLORS = {
    "LinearRegression": "#4C78A8",
    "ElasticNet": "#F58518",
    "RandomForest": "#54A24B",
    "XGBoost": "#E45756",
    "CatBoost": "#72B7B2",
}


def gost_layout(fig, title, y_title, x_title="Широтная зона"):
    """Настраивает оформление Plotly-графика под требования статьи."""
    fig.update_layout(
        title={"text": title, "x": 0.5, "xanchor": "center"},
        font={"family": "Times New Roman", "size": 14, "color": "black"},
        plot_bgcolor="white",
        paper_bgcolor="white",
        width=1100,
        height=650,
        margin={"l": 85, "r": 35, "t": 80, "b": 85},
        legend={
            "orientation": "h",
            "yanchor": "bottom",
            "y": -0.28,
            "xanchor": "center",
            "x": 0.5,
            "font": {"family": "Times New Roman", "size": 14},
        },
    )
    fig.update_xaxes(
        title={"text": x_title, "font": {"family": "Times New Roman", "size": 14}},
        showgrid=True,
        gridcolor="#D9D9D9",
        griddash="dot",
        zeroline=False,
        linecolor="black",
        mirror=True,
        ticks="outside",
        tickfont={"family": "Times New Roman", "size": 14},
    )
    fig.update_yaxes(
        title={"text": y_title, "font": {"family": "Times New Roman", "size": 14}},
        showgrid=True,
        gridcolor="#D9D9D9",
        griddash="dot",
        zeroline=False,
        linecolor="black",
        mirror=True,
        ticks="outside",
        tickfont={"family": "Times New Roman", "size": 14},
    )
    return fig


def save_plotly(fig, stem):
    """Сохраняет Plotly-график в HTML и PNG."""
    html = OUT_DIR / f"{stem}.html"
    png = OUT_DIR / f"{stem}.png"
    fig.write_html(html, include_plotlyjs="cdn")
    fig.write_image(png, scale=2)
    return png


def build_figures(station_model, zone_model):
    """Строит графики для статьи по широтным зонам."""
    x_labels = [ZONE_LABELS[z] for z in ZONE_ORDER]

    fig_rmse = go.Figure()
    for model in MODEL_ORDER:
        m = zone_model[zone_model["model"] == model].set_index("latitude_zone")
        fig_rmse.add_trace(
            go.Scatter(
                x=x_labels,
                y=[m.loc[z, "rmse_mean"] for z in ZONE_ORDER],
                mode="lines+markers",
                name=model,
                line={"width": 2, "color": COLORS[model]},
                marker={"size": 9, "symbol": "circle", "line": {"width": 1, "color": "black"}},
            )
        )
    gost_layout(
        fig_rmse,
        "Средний RMSE моделей по широтным зонам",
        "RMSE foF2, МГц",
    )
    rmse_png = save_plotly(fig_rmse, "gost_plotly_rmse_by_zone_model")

    fig_mae = go.Figure()
    for model in MODEL_ORDER:
        m = zone_model[zone_model["model"] == model].set_index("latitude_zone")
        fig_mae.add_trace(
            go.Scatter(
                x=x_labels,
                y=[m.loc[z, "mae_mean"] for z in ZONE_ORDER],
                mode="lines+markers",
                name=model,
                line={"width": 2, "color": COLORS[model]},
                marker={"size": 9, "symbol": "square", "line": {"width": 1, "color": "black"}},
            )
        )
    gost_layout(
        fig_mae,
        "Средний MAE моделей по широтным зонам",
        "MAE foF2, МГц",
    )
    mae_png = save_plotly(fig_mae, "gost_plotly_mae_by_zone_model")

    fig_rmse_nl = go.Figure()
    for model in ["RandomForest", "XGBoost", "CatBoost"]:
        m = zone_model[zone_model["model"] == model].set_index("latitude_zone")
        fig_rmse_nl.add_trace(
            go.Scatter(
                x=x_labels,
                y=[m.loc[z, "rmse_mean"] for z in ZONE_ORDER],
                mode="lines+markers",
                name=model,
                line={"width": 2, "color": COLORS[model]},
                marker={"size": 10, "symbol": "circle", "line": {"width": 1, "color": "black"}},
            )
        )
    gost_layout(
        fig_rmse_nl,
        "Средний RMSE нелинейных моделей по широтным зонам",
        "RMSE foF2, МГц",
    )
    rmse_nl_png = save_plotly(fig_rmse_nl, "gost_plotly_rmse_nonlinear_by_zone")

    fig_mae_nl = go.Figure()
    for model in ["RandomForest", "XGBoost", "CatBoost"]:
        m = zone_model[zone_model["model"] == model].set_index("latitude_zone")
        fig_mae_nl.add_trace(
            go.Scatter(
                x=x_labels,
                y=[m.loc[z, "mae_mean"] for z in ZONE_ORDER],
                mode="lines+markers",
                name=model,
                line={"width": 2, "color": COLORS[model]},
                marker={"size": 10, "symbol": "square", "line": {"width": 1, "color": "black"}},
            )
        )
    gost_layout(
        fig_mae_nl,
        "Средний MAE нелинейных моделей по широтным зонам",
        "MAE foF2, МГц",
    )
    mae_nl_png = save_plotly(fig_mae_nl, "gost_plotly_mae_nonlinear_by_zone")

    fig_corr = go.Figure()
    for model in ["RandomForest", "XGBoost", "CatBoost"]:
        m = zone_model[zone_model["model"] == model].set_index("latitude_zone")
        fig_corr.add_trace(
            go.Scatter(
                x=x_labels,
                y=[m.loc[z, "corr_mean"] for z in ZONE_ORDER],
                mode="lines+markers",
                name=model,
                line={"width": 2, "color": COLORS[model]},
                marker={"size": 9, "symbol": "diamond", "line": {"width": 1, "color": "black"}},
            )
        )
    gost_layout(
        fig_corr,
        "Корреляция прогноза и наблюдений по широтным зонам",
        "Коэффициент корреляции",
    )
    corr_png = save_plotly(fig_corr, "gost_plotly_corr_by_zone_model")

    fig_box = go.Figure()
    for zone in ZONE_ORDER:
        values = station_model.loc[station_model["latitude_zone"] == zone, "rmse"]
        fig_box.add_trace(
            go.Box(
                y=values,
                name=ZONE_LABELS[zone],
                boxpoints="all",
                jitter=0.35,
                pointpos=0,
                marker={"size": 6, "line": {"width": 0.8, "color": "black"}},
                line={"width": 1.5},
            )
        )
    gost_layout(
        fig_box,
        "Распределение RMSE по широтным зонам",
        "RMSE foF2, МГц",
        x_title="",
    )
    fig_box.update_layout(showlegend=False)
    box_png = save_plotly(fig_box, "gost_plotly_rmse_distribution_by_zone")
    return [rmse_png, mae_png, rmse_nl_png, mae_nl_png, corr_png, box_png]


def set_document_style(doc):
    """Настраивает базовый стиль Word-документа."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Добавляет абзац с единым оформлением."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return p


def add_center(doc, text, bold=False):
    """Добавляет центрированный абзац."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return p


def add_table(doc, data, columns):
    """Добавляет таблицу в Word-документ."""
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, col in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = col
    for _, row in data.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = str(row[col])
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def build_article(figures, zone_summary, zone_model, best_counts, metadata):
    """Собирает Word-документ статьи из таблиц и рисунков."""
    doc = Document()
    set_document_style(doc)

    add_paragraph(doc, "УДК 004.85:550.388", WD_ALIGN_PARAGRAPH.LEFT)
    add_center(doc, "Короткова Надежда Александровна")
    add_center(doc, "Научный руководитель: Конкин Никита Александрович, ст. преп. кафедры Радиотехники и связи")
    add_center(doc, "ФГБОУ ВО «Поволжский государственный технологический университет», г. Йошкар-Ола")
    add_center(
        doc,
        "Влияние широтной зоны ионосферной станции на точность 24-часового ML-прогноза критической частоты слоя F2",
        bold=True,
    )

    add_paragraph(
        doc,
        "Аннотация. В работе исследуется влияние широтного положения ионосферных станций на точность 24-часового прогноза критической частоты слоя F2 (foF2). Для 35 станций сети GIRO выполнено сравнение моделей Linear Regression, ElasticNet, Random Forest, XGBoost и CatBoost. Качество прогноза оценивалось по метрикам MAE, RMSE, R² и коэффициенту корреляции. Станции были распределены по широтным зонам: низкие широты, южные средние широты, северные средние широты и северные высокие широты. Показано, что наиболее точный прогноз достигается в средних широтах, а увеличение ошибки в низких и высоких широтах связано с более сложной ионосферной динамикой.",
    )
    add_paragraph(
        doc,
        "Ключевые слова: foF2, ионосфера, критическая частота, машинное обучение, широтная зона, Random Forest, XGBoost, CatBoost.",
        WD_ALIGN_PARAGRAPH.LEFT,
    )

    add_paragraph(
        doc,
        "Критическая частота слоя F2 (foF2) является одним из основных параметров, характеризующих состояние ионосферы и условия распространения коротких радиоволн. Значение foF2 определяет максимальную частоту радиосигнала, отражающегося от слоя F2 при вертикальном падении, поэтому ее прогнозирование важно для задач радиосвязи, навигации и мониторинга космической погоды.",
    )
    add_paragraph(
        doc,
        "Точность прогноза foF2 существенно зависит от геофизических условий. В низких широтах на динамику электронной концентрации влияет экваториальная ионосферная аномалия, в высоких широтах - авроральные процессы и геомагнитные возмущения, а в средних широтах изменения часто имеют более регулярный суточный и сезонный характер. Поэтому анализ качества прогноза по широтным зонам позволяет определить области, в которых модели машинного обучения наиболее устойчивы.",
    )
    add_paragraph(
        doc,
        "Цель работы - оценить влияние широтного положения ионосферной станции на точность 24-часового ML-прогноза foF2.",
    )
    add_paragraph(
        doc,
        "Для достижения цели были решены следующие задачи: проведены расчеты нескольких моделей машинного обучения для набора ионосферных станций; получены метрики качества для каждой станции; станции распределены по широтным зонам; выполнено сравнение MAE, RMSE, R² и корреляции; определены наиболее устойчивые модели для каждой широтной зоны.",
    )
    add_paragraph(
        doc,
        "В исследовании использованы данные ионосферных станций за 2024-2025 годы с временным шагом 15 минут. Прогноз строился на горизонт 24 часа. В качестве моделей применялись Linear Regression, ElasticNet, Random Forest, XGBoost и CatBoost. Для оценки качества использовались MAE, RMSE, R², коэффициент корреляции и MAPE. Эксперимент проводился в режиме walk-forward daily, что позволяет имитировать последовательное обновление прогноза во времени.",
    )

    station_zone = metadata.groupby("latitude_zone").agg(
        stations=("station", "nunique"), lat_min=("latitude", "min"), lat_max=("latitude", "max")
    ).reindex(ZONE_ORDER).reset_index()
    station_zone["latitude_zone"] = station_zone["latitude_zone"].map(ZONE_LABELS)
    station_zone["lat_min"] = station_zone["lat_min"].map(lambda x: f"{x:.2f}")
    station_zone["lat_max"] = station_zone["lat_max"].map(lambda x: f"{x:.2f}")
    add_center(doc, "Таблица 1 - Распределение станций по широтным зонам")
    add_table(doc, station_zone, ["latitude_zone", "stations", "lat_min", "lat_max"])

    add_paragraph(
        doc,
        "Результаты сравнения показывают, что при рассмотрении нелинейных моделей минимальные средние ошибки получены в северных и южных средних широтах. Для зоны N_mid средний MAE составил 0,808 МГц, RMSE - 1,036 МГц, средняя корреляция - 0,882. Для зоны S_mid значения близки: MAE - 0,809 МГц, RMSE - 1,000 МГц, корреляция - 0,900. В низких широтах RMSE увеличивается до 1,508 МГц, что указывает на более сложную структуру временной изменчивости foF2.",
    )
    add_center(doc, "Таблица 2 - Средние метрики по широтным зонам для нелинейных моделей")
    table2 = zone_summary.reset_index()
    table2["latitude_zone"] = table2["latitude_zone"].map(ZONE_LABELS)
    for col in ["mae_mean", "rmse_mean", "r2_mean", "corr_mean"]:
        table2[col] = table2[col].map(lambda x: f"{x:.3f}")
    add_table(doc, table2[["latitude_zone", "stations", "mae_mean", "rmse_mean", "r2_mean", "corr_mean"]], ["latitude_zone", "stations", "mae_mean", "rmse_mean", "r2_mean", "corr_mean"])

    captions = [
        "Рис. 1. Средний RMSE моделей машинного обучения по широтным зонам.",
        "Рис. 2. Средний MAE моделей машинного обучения по широтным зонам.",
        "Рис. 3. Средний RMSE нелинейных моделей по широтным зонам.",
        "Рис. 4. Средний MAE нелинейных моделей по широтным зонам.",
        "Рис. 5. Корреляция прогноза и наблюдаемых значений foF2 по широтным зонам.",
        "Рис. 6. Распределение RMSE по широтным зонам.",
    ]
    for path, caption in zip(figures, captions):
        doc.add_picture(str(path), width=Cm(15.5))
        add_center(doc, caption)

    add_paragraph(
        doc,
        "Сравнение моделей внутри широтных зон показывает преимущество ансамблевых алгоритмов. Наиболее часто лучшей по RMSE моделью оказывалась CatBoost: она была лучшей для 8 станций в низких широтах, 3 станций в южных средних широтах, 10 станций в северных средних широтах и 2 станций в северных высоких широтах. Random Forest также показал устойчивые результаты, особенно в северных средних широтах, где он оказался лучшим для 8 станций.",
    )
    add_paragraph(
        doc,
        "Линейная регрессия и ElasticNet в большинстве зон уступают ансамблевым методам. Особенно заметно снижение качества по R², что отражает нелинейный характер связи foF2 с локальным временем, предысторией состояния ионосферы и геомагнитными параметрами. Поэтому для практического 24-часового прогноза foF2 предпочтительно использовать CatBoost, Random Forest или XGBoost.",
    )
    add_paragraph(
        doc,
        "Физическая интерпретация полученных различий связана с особенностями широтной структуры ионосферы. В средних широтах суточная и сезонная изменчивость foF2 выражена достаточно регулярно, что облегчает построение прогноза на основе лаговых и временных признаков. В низких широтах существенную роль играют экваториальная аномалия, плазменные неоднородности и быстрые вечерние изменения электронной концентрации. В высоких широтах качество прогноза ограничивается воздействием авроральных процессов, магнитосферно-ионосферного взаимодействия и геомагнитных возмущений. При этом зона северных высоких широт представлена только двумя станциями, поэтому выводы по ней следует рассматривать как предварительные.",
    )
    add_paragraph(
        doc,
        "Заключение. В работе выполнен анализ влияния широтной зоны станции на точность 24-часового ML-прогноза foF2. Наиболее точные результаты получены в средних широтах: средний RMSE нелинейных моделей составляет около 1,0-1,04 МГц при корреляции 0,88-0,90. В низких широтах ошибка возрастает до 1,51 МГц, что связано с более выраженной пространственно-временной неоднородностью ионосферы. Ансамблевые модели CatBoost, Random Forest и XGBoost демонстрируют наибольшую устойчивость, при этом CatBoost чаще всего оказывается лучшей моделью по RMSE для отдельных станций.",
    )
    add_paragraph(doc, "Список литературы:", WD_ALIGN_PARAGRAPH.LEFT)
    refs = [
        "1. Bilitza D. International Reference Ionosphere 2020 / D. Bilitza, D. Altadill, V. Truhlik, D. Buresova, I. Galkin // Reviews of Geophysics. - 2022. - Vol. 60, Iss. 4. - DOI: 10.1029/2021RG000753.",
        "2. Razin B. V. Global Ionospheric Radio Observatory (GIRO) / B. V. Razin, I. A. Galkin // Earth, Planets and Space. - 2011. - Vol. 63. - P. 377-381. - DOI: 10.5047/eps.2011.03.001.",
        "3. Mao S. A Review of Machine Learning-Based Ionospheric Spatial and Temporal Modeling / S. Mao, M. Hernández-Pajares, B. Soja // Journal of Geophysical Research: Space Physics. - 2024. - DOI: 10.1029/2024jh000555.",
        "4. Журавлёв С. В. Краткосрочный прогноз критической частоты ионосферного слоя F2 / С. В. Журавлёв, Н. Г. Котонаева, А. В. Михайлов, В. В. Михайлов, К. Г. Цыбуля // Результаты испытания новых и усовершенствованных технологий, моделей и методов гидрометеорологических прогнозов. - 2022. - № 49. - С. 131-146.",
    ]
    for ref in refs:
        add_paragraph(doc, ref)

    path = OUT_DIR / "Статья_foF2_широтные_зоны_ГОСТ.docx"
    doc.save(path)
    return path


def main():
    """Запускает основной сценарий файла."""
    station_model = pd.read_csv(REPORT_DIR / "station_latitude_zone_model_errors.csv")
    zone_model = pd.read_csv(REPORT_DIR / "zone_model_summary.csv")
    zone_summary = pd.read_csv(REPORT_DIR / "zone_summary_nonlinear_models.csv", index_col=0)
    best_counts = pd.read_csv(REPORT_DIR / "best_model_counts_by_zone.csv", index_col=0)
    metadata = pd.read_csv(ROOT / "reports" / "fof2_fast_ml_24h_station_metadata.csv")
    figures = build_figures(station_model, zone_model)
    article = build_article(figures, zone_summary, zone_model, best_counts, metadata)
    print(article)
    for fig in figures:
        print(fig)


if __name__ == "__main__":
    main()
